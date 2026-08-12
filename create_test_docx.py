import os
import docx
from metadata_engine import WordMetadataEngine

test_dir = os.path.abspath('test_documents')
os.makedirs(test_dir, exist_ok=True)

test_files = [
    {
        'filename': '测试文档_1_项目报告.docx',
        'title': '2026年半年度项目规划与进度汇报',
        'author': 'python-docx',
        'last_modified_by': '张三',
        'company': '星空科技集团',
        'comments': '本文档包含项目研发核心规划与保密数据。',
        'total_editing_time': '120',
        'revision': '3',
        'created_time': '2025-06-15T09:30:00Z',
        'modified_time': '2025-06-20T16:45:00Z',
    },
    {
        'filename': '测试文档_2_财务报表.docx',
        'title': '2025第四季度财务审计与预算分析',
        'author': 'python-docx',
        'last_modified_by': '李四',
        'company': '未来金融股份有限公司',
        'comments': '内部审阅件，请勿外传。',
        'total_editing_time': '45',
        'revision': '1',
        'created_time': '2025-10-01T10:00:00Z',
        'modified_time': '2025-10-05T11:20:00Z',
    },
    {
        'filename': '测试文档_3_技术架构.docx',
        'title': '云原生高并发架构设计规范书',
        'author': 'python-docx',
        'last_modified_by': '王五',
        'company': '极客软件研发中心',
        'comments': '技术架构组统一规范文档。',
        'total_editing_time': '310',
        'revision': '12',
        'created_time': '2025-03-10T14:15:00Z',
        'modified_time': '2025-04-01T18:00:00Z',
    },
    {
        'filename': '测试文档_4_会议纪要.docx',
        'title': '产品体验迭代研讨会会议记录',
        'author': 'python-docx',
        'last_modified_by': '赵六',
        'company': '体验设计实验室',
        'comments': '会议要点及后续待办事项跟进表。',
        'total_editing_time': '25',
        'revision': '2',
        'created_time': '2025-11-12T16:00:00Z',
        'modified_time': '2025-11-12T17:30:00Z',
    },
    {
        'filename': '测试文档_5_脱敏测试.docx',
        'title': '个人信息敏感数据测试样本',
        'author': 'python-docx',
        'last_modified_by': '孙七',
        'company': '安全合规部门',
        'comments': '专门用于测试一键隐私脱敏功能。',
        'total_editing_time': '90',
        'revision': '4',
        'created_time': '2025-08-20T08:00:00Z',
        'modified_time': '2025-08-21T10:10:00Z',
    }
]

print(f"开始生成测试文档至目录: {test_dir}\n" + "-"*50)

for idx, data in enumerate(test_files, 1):
    file_path = os.path.join(test_dir, data['filename'])
    
    # 1. Create document
    doc = docx.Document()
    doc.add_heading(data['title'], 0)
    doc.add_paragraph(f"测试文档编号: {idx}")
    doc.add_paragraph(f"文档用途: 用于测试 Word 属性编辑器的单文件精修与批量处理功能。")
    doc.add_paragraph(f"备注说明: {data['comments']}")
    doc.save(file_path)
    
    # 2. Write metadata
    meta = WordMetadataEngine.read_metadata(file_path)
    meta['author'] = data['author']
    meta['last_modified_by'] = data['last_modified_by']
    meta['company'] = data['company']
    meta['title'] = data['title']
    meta['comments'] = data['comments']
    meta['total_editing_time'] = data['total_editing_time']
    meta['revision'] = data['revision']
    meta['created_time'] = data['created_time']
    meta['modified_time'] = data['modified_time']
    
    WordMetadataEngine.write_metadata(file_path, meta, sync_fs_time=True)
    print(f"[{idx}/5] 成功生成: {data['filename']}")

print("-" * 50 + "\n所有测试文档已重新生成完毕！")
